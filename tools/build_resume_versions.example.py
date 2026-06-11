from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = str(Path.home() / "Desktop" / "job application" / "resume" / "standard resume")

CANDIDATE_NAME = "Your Name"
CONTACT = "Your Name | your.email@example.com | linkedin.com/in/your-profile | City, ST"

BASE_EXPERIENCE = [
    {
        "role": "Full Stack Developer",
        "company": "Your Company",
        "location": "City, ST",
        "dates": "Jan 2022 - Present",
        "bullets": {
            "dotnet": [
                "Built production web applications using C#, .NET, REST APIs, SQL, and modern frontend frameworks.",
                "Improved API performance, data quality, and release reliability through testing and monitoring.",
            ],
            "cloud": [
                "Deployed and supported cloud applications using managed app services, databases, and CI/CD pipelines.",
            ],
            "ai": [
                "Integrated AI-assisted workflows, prompt engineering, or retrieval-based features where relevant.",
            ],
            "integration": [
                "Integrated third-party APIs and internal services for business workflows.",
            ],
        },
    },
    {
        "role": "Software Engineer",
        "company": "Previous Company",
        "location": "City, ST",
        "dates": "Jan 2020 - Dec 2021",
        "bullets": {
            "dotnet": [
                "Developed and maintained enterprise features across backend services, frontend UI, and relational databases.",
                "Collaborated with product, QA, and operations teams in Agile delivery cycles.",
            ],
            "cloud": [
                "Supported environment configuration, deployments, and production troubleshooting.",
            ],
        },
    },
]

PROJECTS = {
    "cloud_task": {
        "name": "Cloud Task Manager",
        "stack": ".NET, JavaScript/TypeScript, SQL, Cloud App Service, CI/CD",
        "description": "Built and deployed a task management app with API endpoints, database persistence, and automated deployment.",
    },
    "stock": {
        "name": "Analytics Dashboard",
        "stack": ".NET, Frontend Framework, SQL, Charting Library",
        "description": "Created dashboard screens, API integrations, and reporting views for operational data.",
    },
    "rag": {
        "name": "AI Search Assistant",
        "stack": "Python, LLM API, Vector Search, REST API",
        "description": "Built a retrieval-based assistant prototype that answers questions from indexed documents.",
    },
}

VERSIONS = [
    {
        "filename": "Candidate_FullStack_NET_Cloud.docx",
        "title": "Full Stack .NET Developer | Cloud Engineer",
        "summary": "Full Stack .NET Developer with experience building APIs, frontend applications, databases, cloud deployments, and CI/CD workflows.",
        "skills": [
            ("Languages", "C#, TypeScript, JavaScript, SQL, Python"),
            ("Backend", ".NET, ASP.NET Core, Web API, Entity Framework"),
            ("Frontend", "Angular, React, HTML5, CSS3"),
            ("Cloud & DevOps", "Azure/AWS/GCP, CI/CD, Docker, GitHub Actions or Azure DevOps"),
            ("Databases", "SQL Server, PostgreSQL, MongoDB"),
            ("Practices", "REST APIs, Microservices, Agile/Scrum, Testing, Monitoring"),
        ],
        "experience_order": ["dotnet", "cloud", "integration", "ai"],
        "projects": ["cloud_task", "stock", "rag"],
    },
    {
        "filename": "Candidate_FullStack_NET_Cloud_AI.docx",
        "title": "Full Stack .NET Developer | Cloud & AI Engineer",
        "summary": "Full Stack .NET Developer with cloud engineering experience and practical AI/automation project exposure.",
        "skills": [
            ("Languages", "C#, Python, TypeScript, JavaScript, SQL"),
            ("Full Stack", ".NET, Web API, Entity Framework, Angular, React"),
            ("AI", "LLM APIs, RAG, Prompt Engineering, Vector Search"),
            ("Cloud & DevOps", "Azure/AWS/GCP, CI/CD, Docker, Cloud Databases"),
            ("Databases", "SQL Server, PostgreSQL, MongoDB"),
            ("Practices", "REST APIs, Microservices, Agile/Scrum, Testing, Monitoring"),
        ],
        "experience_order": ["dotnet", "ai", "cloud", "integration"],
        "projects": ["rag", "cloud_task", "stock"],
    },
    {
        "filename": "Candidate_AI_Cloud_Engineer.docx",
        "title": "AI & Cloud Engineer | Automation | Data Pipelines",
        "summary": "AI and cloud engineer with software delivery experience across APIs, data workflows, automation, and analytics.",
        "skills": [
            ("AI", "LLM APIs, RAG, Prompt Engineering, Vector Search"),
            ("Data", "Python, SQL, ETL, Analytics, Dashboards"),
            ("Cloud", "Azure/AWS/GCP App Services, Cloud Databases, Storage, Monitoring"),
            ("Backend", ".NET, FastAPI, REST APIs"),
            ("Frontend", "Angular, React, TypeScript"),
            ("DevOps", "CI/CD, Docker, Git, Infrastructure as Code"),
        ],
        "experience_order": ["ai", "cloud", "dotnet", "integration"],
        "projects": ["rag", "cloud_task", "stock"],
    },
]


def add_text(paragraph, text, bold=False, size=10, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    add_text(paragraph, text.upper(), bold=True, size=10.5, color=RGBColor(31, 78, 121))


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(1.5)
    add_text(paragraph, text, size=9.3)


def experience_bullets(exp, order):
    seen = set()
    bullets = []
    for group in order:
        for item in exp.get("bullets", {}).get(group, []):
            if item not in seen:
                seen.add(item)
                bullets.append(item)
    for group_items in exp.get("bullets", {}).values():
        for item in group_items:
            if item not in seen:
                seen.add(item)
                bullets.append(item)
    return bullets[:5]


def build(version):
    out_dir = Path(OUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)

    name = CONTACT.split("|", 1)[0].strip() or "Candidate Name"
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    add_text(paragraph, name, bold=True, size=16, color=RGBColor(31, 78, 121))

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    add_text(paragraph, CONTACT, size=9)

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    add_text(paragraph, version["title"], bold=True, size=10.5)

    heading(doc, "Professional Summary")
    doc.add_paragraph(version["summary"])

    heading(doc, "Technical Skills")
    for label, values in version.get("skills", []):
        paragraph = doc.add_paragraph()
        add_text(paragraph, f"{label}: ", bold=True, size=9.5)
        add_text(paragraph, values, size=9.5)

    heading(doc, "Professional Experience")
    for exp in BASE_EXPERIENCE:
        paragraph = doc.add_paragraph()
        add_text(paragraph, f"{exp['role']} | {exp['company']}", bold=True, size=10)
        add_text(paragraph, f" | {exp['location']} | {exp['dates']}", size=9.2)
        for item in experience_bullets(exp, version.get("experience_order", [])):
            bullet(doc, item)

    heading(doc, "Projects")
    for key in version.get("projects", []):
        project = PROJECTS.get(key)
        if not project:
            continue
        paragraph = doc.add_paragraph()
        add_text(paragraph, project["name"], bold=True, size=9.8)
        add_text(paragraph, f" | {project['stack']}", size=9.2)
        bullet(doc, project["description"])

    path = out_dir / version["filename"]
    doc.save(path)
    return str(path)


if __name__ == "__main__":
    for item in VERSIONS:
        print(build(item))
