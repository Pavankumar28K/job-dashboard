import json
import os
import re
import shutil
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
try:
    import openai
except ImportError:
    openai = None


APP_DIR = Path(__file__).resolve().parents[1]


def configured_path(env_name, fallback):
    return Path(os.environ.get(env_name, fallback)).expanduser().resolve()


DATA_PATH = configured_path("DASHBOARD_JOBS", APP_DIR / "data" / "jobs.json")
JOB_ROOT = configured_path("JOB_APP_ROOT", Path.home() / "Desktop" / "job application")
TAILORED_DIR = JOB_ROOT / "resume" / "AI tailored resume"
COVER_DIR = JOB_ROOT / "cover letters"
JD_DIR = JOB_ROOT / "job descriptions"
CONFIG_PATH = APP_DIR / "config.json"


def safe_name(value, limit=70):
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", value or "").strip("_")
    return cleaned[:limit] or "Job"


def load_config():
    try:
        return json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def base_resume_path():
    configured = os.environ.get("BASE_RESUME_PATH") or load_config().get("baseResumePath") or ""
    if not configured:
        raise SystemExit("Base resume is not configured. Add a .docx full path in Settings first.")
    path = Path(configured).expanduser().resolve()
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise SystemExit(f"Base resume must be an existing .docx file: {path}")
    return path


def load_job(job_id):
    jobs = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    for job in jobs:
        if job["id"] == job_id:
            return job
    raise SystemExit(f"Job not found: {job_id}")


def document_lines(path):
    doc = Document(path)
    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.extend(line.strip() for line in text.splitlines() if line.strip())
    return lines


def document_blob(path):
    return "\n".join(document_lines(path)).lower()


def candidate_name(path):
    lines = document_lines(path)
    return lines[0] if lines else "Candidate"


def contact_line(path):
    lines = document_lines(path)
    if len(lines) > 1:
        return lines[1]
    return ""


def candidate_safe_name(path):
    return safe_name(candidate_name(path), 40)


def text_blob(job):
    return " ".join(
        str(job.get(key, ""))
        for key in ["company", "role", "source", "location", "employmentType", "notes", "jd", "workAuthRisk"]
    ).lower()


def configured_skill_hits(job):
    config = load_config()
    skills = [*(config.get("mustHaveSkills") or []), *(config.get("niceToHaveSkills") or [])]
    blob = text_blob(job)
    hits = []
    for skill in skills:
        skill_text = str(skill).strip()
        if skill_text and skill_text.lower() in blob and skill_text not in hits:
            hits.append(skill_text)
    return hits


def keyword_hits(job):
    configured = configured_skill_hits(job)
    generic_terms = [
        "Python", "Java", "JavaScript", "TypeScript", "React", "Angular", "Node.js", "SQL",
        "PostgreSQL", "MySQL", "MongoDB", "AWS", "Azure", "GCP", "Docker", "Kubernetes",
        "Terraform", "CI/CD", "REST", "API", "Machine Learning", "Deep Learning", "NLP",
        "LLM", "OpenAI", "LangChain", "Spark", "ETL", "Power BI", "Tableau", "Agile",
        "Scrum", "Git", "FastAPI", "Pydantic", "Django", "Flask", ".NET", "C#", "Spring Boot",
        "Kafka", "Event-Driven", "Microservices", "Automation", "Testing", "Monitoring",
    ]
    blob = text_blob(job)
    hits = configured[:]
    for term in generic_terms:
        if term not in hits and term.lower() in blob:
            hits.append(term)
    return hits[:12]


def resume_supported_keywords(source, job):
    resume_blob = document_blob(source)
    supported = []
    for keyword in keyword_hits(job):
        if keyword.lower() in resume_blob:
            supported.append(keyword)
    return supported


SKILL_LABELS = {
    "API": "REST APIs",
    "REST": "REST APIs",
    "Event-Driven": "event-driven architecture",
    "Machine Learning": "machine learning",
    "CI/CD": "CI/CD pipelines",
}


def display_skills(skills):
    labels = []
    for skill in skills:
        label = SKILL_LABELS.get(skill, skill)
        if label not in labels:
            labels.append(label)
    if len(labels) <= 1:
        return labels[0] if labels else ""
    if len(labels) == 2:
        return f"{labels[0]} and {labels[1]}"
    return ", ".join(labels[:-1]) + f", and {labels[-1]}"


def skills_for_group(supported, terms, fallback_count=3, allow_fallback=False):
    matched = [skill for skill in supported if skill in terms]
    if matched:
        return display_skills(matched[:4])
    if not allow_fallback:
        return ""
    return display_skills(supported[:fallback_count])


def jd_functional_bullets(job, supported):
    blob = text_blob(job)
    bullets = []

    focus_areas = [
        (
            r"\b(api|backend|server|service|microservice|fastapi|spring boot|\.net|pydantic|integration)\b",
            ["REST", "API", "FastAPI", "Pydantic", "Django", "Flask", "Spring Boot", ".NET", "C#", "Java", "Python", "Microservices"],
            "Developed and maintained backend services, REST APIs, and system integrations using {skills} to support scalable production applications.",
        ),
        (
            r"\b(event|queue|stream|kafka|workflow|throughput|distributed|platform|ledger|transaction)\b",
            ["Kafka", "Event-Driven", "Python", "Java", "AWS", "Azure", "Docker", "Kubernetes", "SQL", "PostgreSQL", "Microservices"],
            "Built event-driven processing workflows and distributed platform components using {skills} to improve reliability, throughput, and data consistency.",
        ),
        (
            r"\b(machine learning|ml\b|ai\b|model|predictive|statistical|nlp|llm|openai|langchain|agent|automation)\b",
            ["Machine Learning", "Deep Learning", "NLP", "LLM", "OpenAI", "LangChain", "Python", "Spark", "ETL", "Automation"],
            "Implemented AI, machine learning, and automation solutions with {skills} to support model integration, predictive insights, and application workflows.",
        ),
        (
            r"\b(data|dataset|analytics|insight|pipeline|etl|mining|warehouse|report|dashboard|metrics)\b",
            ["SQL", "PostgreSQL", "MySQL", "MongoDB", "Python", "Spark", "ETL", "Power BI", "Tableau", "Machine Learning"],
            "Analyzed datasets, data pipelines, and reporting workflows using {skills} to deliver actionable insights and support data-driven decisions.",
        ),
        (
            r"\b(cloud|deploy|deployment|devops|ci/cd|container|docker|kubernetes|terraform|aws|azure|gcp|release)\b",
            ["AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "CI/CD", "Git"],
            "Deployed and supported cloud, container, and CI/CD release workflows using {skills} to improve production stability and delivery speed.",
        ),
        (
            r"\b(frontend|front-end|ui\b|ux\b|web app|single page|component|responsive|react|angular)\b",
            ["React", "Angular", "JavaScript", "TypeScript", "HTML", "CSS"],
            "Developed responsive web application features, UI components, and frontend integrations using {skills} to improve user experience and maintainability.",
        ),
        (
            r"\b(test|testing|quality|qa|validation|debug|troubleshoot|monitor|observability|reliability|incident|production support)\b",
            ["Testing", "Monitoring", "Git", "CI/CD", "Python", "JavaScript", "Java", "C#"],
            "Tested, debugged, monitored, and supported production systems using {skills} to improve reliability, quality, and issue resolution.",
        ),
        (
            r"\b(security|secure|auth|authorization|authentication|compliance|audit|audited|risk|privacy)\b",
            ["API", "REST", ".NET", "C#", "Java", "Python", "SQL", "Azure", "AWS"],
            "Implemented secure application, API, and data workflows using {skills} to support authentication, compliance, audit readiness, and risk reduction.",
        ),
        (
            r"\b(product|stakeholder|cross-functional|collaborat|present|recommendation|requirement|agile|scrum)\b",
            ["Agile", "Scrum", "Git", "Power BI", "Tableau", "SQL", "Python", "Machine Learning"],
            "Collaborated with product, engineering, and cross-functional stakeholders using {skills} to gather requirements, deliver Agile increments, and present recommendations.",
        ),
    ]

    used_text = set()
    for pattern, terms, template in focus_areas:
        if not re.search(pattern, blob, re.I):
            continue
        skills = skills_for_group(supported, terms)
        if not skills:
            continue
        bullet = template.format(skills=skills)
        if bullet not in used_text:
            bullets.append(bullet)
            used_text.add(bullet)
        if len(bullets) >= 4:
            break

    return bullets


def relevant_experience_bullets(source, job):
    supported = resume_supported_keywords(source, job)
    if not supported:
        return []

    bullets = jd_functional_bullets(job, supported)

    skill_groups = [
        (["REST", "API", "FastAPI", "Pydantic", "Django", "Flask", "Spring Boot", ".NET", "C#", "Java", "Python"], "Implemented application services, API endpoints, and backend logic using {skills} to support maintainable feature delivery."),
        (["React", "Angular", "JavaScript", "TypeScript", "HTML", "CSS"], "Developed frontend components and user-facing workflows using {skills}."),
        (["SQL", "PostgreSQL", "MySQL", "MongoDB"], "Designed, queried, and maintained application data flows using {skills}."),
        (["AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "CI/CD"], "Managed deployment, environment, and release workflows using {skills}."),
        (["Machine Learning", "Deep Learning", "NLP", "LLM", "OpenAI", "LangChain", "Spark", "ETL"], "Implemented data, AI, and automation workflows using {skills} to support delivery outcomes."),
        (["Agile", "Scrum", "Git"], "Participated in Agile delivery cycles using {skills} to translate requirements into shipped improvements."),
    ]

    used = set()
    for terms, template in skill_groups:
        matched = [skill for skill in supported if skill in terms and skill not in used]
        if matched:
            bullet = template.format(skills=display_skills(matched[:4]))
            if bullet not in bullets:
                bullets.append(bullet)
            used.update(matched)
        if len(bullets) >= 4:
            break

    if len(bullets) < 3 and len(supported) > 4:
        bullets.append(f"Applied role-relevant technical skills, including {display_skills(supported[4:8])}, to support requirements, delivery, and production support needs described in the job posting.")

    return bullets[:4]


EXPERIENCE_HEADINGS = {
    "experience",
    "professional experience",
    "work experience",
    "employment history",
    "professional background",
}

STOP_HEADINGS = {
    "education",
    "skills",
    "technical skills",
    "projects",
    "project experience",
    "certifications",
    "certification",
    "licenses",
    "awards",
    "summary",
    "professional summary",
}


def clean_heading(text):
    return re.sub(r"[^a-z0-9+#. ]+", " ", text.lower()).strip()


def compact_words(text):
    return [word for word in re.split(r"\s+", text.strip()) if word]


def is_heading_like(paragraph):
    text = paragraph.text.strip()
    if not text:
        return False
    words = compact_words(text)
    style_name = (paragraph.style.name or "").lower()
    if "heading" in style_name:
        return True
    if len(words) <= 5 and text.upper() == text and any(char.isalpha() for char in text):
        return True
    return False


def is_experience_heading(paragraph):
    text = clean_heading(paragraph.text)
    return text in EXPERIENCE_HEADINGS and len(compact_words(paragraph.text)) <= 4


def is_stop_heading(paragraph):
    text = clean_heading(paragraph.text)
    return text in STOP_HEADINGS and len(compact_words(paragraph.text)) <= 4


def is_bullet_like(paragraph):
    text = paragraph.text.strip()
    style_name = (paragraph.style.name or "").lower()
    if "bullet" in style_name or "list" in style_name:
        return True
    if text.startswith(("-", "•", "*", "◦", "▪")):
        return True
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def dateish_text(text):
    return re.search(
        r"\b(19|20)\d{2}\b|\bpresent\b|\bjan\b|\bfeb\b|\bmar\b|\bapr\b|\bmay\b|\bjun\b|\bjul\b|\baug\b|\bsep\b|\boct\b|\bnov\b|\bdec\b",
        text,
        re.I,
    )


def roleish_text(text):
    return re.search(
        r"\b(engineer|developer|analyst|consultant|architect|specialist|lead|intern|manager|scientist|administrator)\b",
        text,
        re.I,
    )


def companyish_text(text):
    return re.search(
        r"\b(capital one|applied materials|inc|llc|corp|corporation|company|systems|solutions|technologies|tech|labs|bank|university|cigna|comcast|wipro|tcs|illumina)\b",
        text,
        re.I,
    )


def separated_text(text):
    return any(token in text for token in ("|", " - ", " – ", " — ", ","))


def is_probable_company_header(paragraph, following):
    text = paragraph.text.strip()
    if not text or is_stop_heading(paragraph) or is_environment_line(paragraph):
        return False
    words = compact_words(text)
    if len(words) > 18:
        return False
    if dateish_text(text) and roleish_text(text):
        return False
    future_text = " ".join(p.text.strip() for p in following[:4])
    has_work_content = len(compact_words(future_text)) >= 6
    if not has_work_content:
        return False
    if dateish_text(text) and not roleish_text(text):
        return True
    if companyish_text(text) and (dateish_text(future_text) or roleish_text(future_text)):
        return True
    if separated_text(text) and (dateish_text(text) or roleish_text(future_text)):
        return True
    return False


def is_environment_line(paragraph):
    return paragraph.text.strip().lower().startswith("environment:")


def is_insertable_experience_line(paragraph):
    text = paragraph.text.strip()
    if not text:
        return False
    if is_stop_heading(paragraph) or is_environment_line(paragraph):
        return False
    if text.lower().rstrip(":") in {"responsibilities", "tools", "technologies"}:
        return False
    if len(compact_words(text)) < 6 and not is_bullet_like(paragraph):
        return False
    return True


def copy_run_format(source_run, target_run):
    if source_run is None:
        return
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if paragraph._p.pPr is not None:
        new_paragraph._p.append(deepcopy(paragraph._p.pPr))
    if style is not None:
        try:
            new_paragraph.style = style
        except Exception:
            pass
    new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after or Pt(3)
    run = new_paragraph.add_run(text)
    copy_run_format(paragraph.runs[0] if paragraph.runs else None, run)
    return new_paragraph


def experience_ranges(doc):
    ranges = []
    start = None
    for index, paragraph in enumerate(doc.paragraphs):
        if is_experience_heading(paragraph):
            start = index + 1
            continue
        if start is not None and is_stop_heading(paragraph):
            if index > start:
                ranges.append((start, index))
            start = None
    if start is not None and start < len(doc.paragraphs):
        ranges.append((start, len(doc.paragraphs)))
    return ranges


def experience_insert_points(doc):
    points = []
    for start, end in experience_ranges(doc):
        paragraphs = doc.paragraphs
        headers = []
        for index in range(start, end):
            following = paragraphs[index + 1 : min(end, index + 7)]
            if is_probable_company_header(paragraphs[index], following):
                headers.append(index)

        if not headers:
            content = [index for index in range(start, end) if is_insertable_experience_line(paragraphs[index])]
            if content:
                points.append(content[-1])
            continue

        for position, header_index in enumerate(headers):
            next_header = headers[position + 1] if position + 1 < len(headers) else end
            block_indices = range(header_index + 1, next_header)
            content = [index for index in block_indices if is_insertable_experience_line(paragraphs[index])]
            if content:
                points.append(content[-1])
    return points


def weave_role_alignment_into_experience(doc, job, source):
    api_key = __import__('os').environ.get("OPENAI_API_KEY", "").strip()
    if api_key and openai:
        return ai_weave_role_alignment(doc, job, api_key)

    bullets = relevant_experience_bullets(source, job)
    if not bullets:
        return 0

    insert_points = experience_insert_points(doc)
    inserted = 0
    for idx, point in enumerate(reversed(insert_points)):
        base_paragraph = doc.paragraphs[point]
        style = base_paragraph.style if is_bullet_like(base_paragraph) else None
        bullet = bullets[idx % len(bullets)]
        base_text = base_paragraph.text.strip()
        prefix = "- " if base_text.startswith("-") else ""
        insert_paragraph_after(base_paragraph, f"{prefix}{bullet}", style=style)
        inserted += 1
    return inserted

import json

def ai_batch_rewrite_resume(doc_data, jd_text, api_key):
    client = openai.OpenAI(api_key=api_key)
    
    prompt = (
        "You are an expert resume writer. I will provide a JSON object containing the current Title, Skills section, and multiple "
        "Experience bullet point blocks from a candidate's resume. I will also provide a Target Job Description.\n"
        "Your task is to rewrite the Title, Skills, and ALL bullet points to perfectly align with the job description.\n\n"
        "RULES:\n"
        "1. **STRICT DOMAIN & COMPANY PRESERVATION**: You are STRICTLY FORBIDDEN from changing the business domain or industry of the candidate's past projects. If the candidate worked at a bank, it MUST remain a banking project. DO NOT inject the Job Description's industry or domain into the candidate's history under any circumstances. You must maintain the original project context.\n"
        "2. **MAINTAIN POINT COUNT**: Keep the exact same number of bullet points per experience block as the original resume.\n"
        "3. **TECHNOLOGY REPLACEMENT**: Replace the technologies in the resume with respect to the JD. If there is any alternate technology already mentioned in the resume that solves the same problem, replace the point with the JD's technology. However, if there are important skills already mentioned in the base resume, do NOT remove them.\n"
        "4. **HIGHLIGHT SKILLS**: You MUST highlight the technical skills and keywords within the bullet points using markdown double asterisks (e.g., **Python**, **AWS**).\n"
        "5. **FORMATTING**: Keep the proper formatting exactly as the base resume.\n"
        "6. Return EXACTLY a valid JSON object matching this structure without extra text:\n"
        "{\n"
        "  \"title\": \"New Job Title\",\n"
        "  \"skills\": [\"Skill 1, Skill 2, Skill 3\", \"Skill 4, Skill 5\"],\n"
        "  \"experiences\": {\n"
        "    \"exp_0\": [\"Rewritten bullet **skill** 1\", \"Rewritten bullet 2\"]\n"
        "  }\n"
        "}\n\n"
        f"--- TARGET JOB DESCRIPTION ---\n{jd_text}\n\n"
        f"--- ORIGINAL RESUME JSON ---\n{json.dumps(doc_data, indent=2)}\n"
    )
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        content = response.choices[0].message.content.strip()
        clean_content = content.replace('```json', '').replace('```', '').strip()
        return json.loads(clean_content)
    except Exception as e:
        print(f"OpenAI batch generation failed: {e}")
        return None

def ai_weave_role_alignment(doc, job, api_key):
    jd_text = text_blob(job)
    
    doc_data = {
        "title": "",
        "skills": [],
        "experiences": {}
    }
    
    # Extract Title
    title_idx = -1
    for i in range(min(5, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text and '|' in text and '@' not in text and 'linkedin' not in text.lower():
            title_idx = i
            break
    if title_idx == -1: title_idx = 1
    doc_data["title"] = doc.paragraphs[title_idx].text.strip()
    
    # Extract Skills
    skills_start, skills_end = -1, -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if text in ['key skills', 'skills', 'technical skills']:
            skills_start = i + 1
            break
    if skills_start != -1:
        for i in range(skills_start, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip().lower()
            if not text or text in ['work experience', 'experience', 'professional experience']:
                skills_end = i
                break
    
    skills_map = []
    if skills_start != -1 and skills_end != -1:
        for i in range(skills_start, skills_end):
            text = doc.paragraphs[i].text.strip()
            if text:
                skills_map.append(i)
                doc_data["skills"].append(text)
                
    # Extract Experiences
    exp_map = {}
    exp_counter = 0
    
    # Assume experience_ranges and is_probable_company_header exist
    for start, end in experience_ranges(doc):
        paragraphs = doc.paragraphs
        headers = []
        for index in range(start, end):
            following = paragraphs[index + 1 : min(end, index + 7)]
            if is_probable_company_header(paragraphs[index], following):
                headers.append(index)
        if not headers: headers = [start - 1]

        for position, header_index in enumerate(headers):
            next_header = headers[position + 1] if position + 1 < len(headers) else end
            block_indices = range(max(start, header_index + 1), next_header)
            
            old_bullets = []
            for idx in block_indices:
                if is_insertable_experience_line(paragraphs[idx]):
                    old_bullets.append(idx)
            
            if old_bullets:
                exp_id = f"exp_{exp_counter}"
                doc_data["experiences"][exp_id] = [paragraphs[idx].text.strip() for idx in old_bullets]
                exp_map[exp_id] = old_bullets
                exp_counter += 1

    print(f"Batching prompt with {len(exp_map)} experiences...")
    rewritten_data = ai_batch_rewrite_resume(doc_data, jd_text, api_key)
    if not rewritten_data:
        return 0
        
    rewritten_count = 0
    
    def apply_text(paragraph, new_text):
        # Clear existing runs safely
        if paragraph.runs:
            for r in paragraph.runs:
                r.text = ""
        
        # Parse **bold** markdown and create runs
        parts = new_text.split('**')
        is_bold = False
        for part in parts:
            if part:
                run = paragraph.add_run(part)
                if is_bold:
                    run.bold = True
            is_bold = not is_bold 

    # Apply Title
    if "title" in rewritten_data and rewritten_data["title"]:
        apply_text(doc.paragraphs[title_idx], rewritten_data["title"])
        rewritten_count += 1
        
    # Apply Skills
    if "skills" in rewritten_data and len(rewritten_data["skills"]) == len(skills_map):
        for idx, new_text in zip(skills_map, rewritten_data["skills"]):
            apply_text(doc.paragraphs[idx], new_text)
            rewritten_count += 1
            
    # Apply Experiences
    if "experiences" in rewritten_data:
        for exp_id, new_bullets in rewritten_data["experiences"].items():
            if exp_id in exp_map:
                old_indices = exp_map[exp_id]
                
                for i in range(min(len(old_indices), len(new_bullets))):
                    apply_text(doc.paragraphs[old_indices[i]], new_bullets[i])
                    rewritten_count += 1
                    
                if len(new_bullets) < len(old_indices):
                    for i in range(len(new_bullets), len(old_indices)):
                        apply_text(doc.paragraphs[old_indices[i]], "")

    return rewritten_count

def paragraph(doc, text, size=10.5, bold=False, color=None, align=None, after=7):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def make_resume(job):
    source = base_resume_path()
    TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    for temp_file in TAILORED_DIR.glob("~$*.docx"):
        try:
            temp_file.unlink()
        except OSError:
            pass

    doc = Document(source)
    weave_role_alignment_into_experience(doc, job, source)
    candidate = candidate_safe_name(source)
    company = job.get("company", "Company")
    role = job.get("role", "Role")
    out = TAILORED_DIR / f"{job['id']}_{safe_name(company, 32)}_{safe_name(role, 40)}_{candidate}.docx"
    upload_copy = TAILORED_DIR / f"UPLOAD_{safe_name(job['id'], 28)}_{candidate}.docx"
    latest_copy = TAILORED_DIR / f"UPLOAD_LATEST_AI_TAILORED_RESUME_{candidate}.docx"
    doc.save(out)
    shutil.copy2(out, upload_copy)
    shutil.copy2(out, latest_copy)
    return upload_copy


def make_cover(job):
    source = base_resume_path()
    COVER_DIR.mkdir(parents=True, exist_ok=True)
    company = job.get("company", "Hiring Team")
    role = job.get("role", "the role")
    keywords = keyword_hits(job)
    name = candidate_name(source)
    contact = contact_line(source)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    paragraph(doc, name, size=16, bold=True, color=RGBColor(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    if contact:
        paragraph(doc, contact, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    alignment = ", ".join(keywords[:7]) if keywords else "the requirements described in the posting"
    lines = [
        date.today().strftime("%B %d, %Y"),
        f"Hiring Team\n{company}",
        f"Re: {role}",
        (
            f"Dear Hiring Team, I am interested in the {role} opportunity with {company}. "
            f"My resume has been tailored around {alignment} for this specific posting."
        ),
        (
            "I would welcome the chance to discuss how my background can support your team. "
            "The attached resume provides the detailed experience, skills, and project history most relevant to this role."
        ),
        f"Sincerely\n{name}",
    ]
    for line in lines:
        paragraph(doc, line)

    out = COVER_DIR / f"{job['id']}_{safe_name(company, 32)}_{safe_name(role, 40)}_Cover_Letter_{candidate_safe_name(source)}.docx"
    doc.save(out)
    return out


def save_jd(job):
    JD_DIR.mkdir(parents=True, exist_ok=True)
    out = JD_DIR / f"{job['id']}_{safe_name(job.get('company'), 32)}_{safe_name(job.get('role'), 40)}_JD.txt"
    content = "\n".join(
        [
            f"Job ID: {job.get('id', '')}",
            f"Company: {job.get('company', '')}",
            f"Role: {job.get('role', '')}",
            f"Source: {job.get('source', '')}",
            f"URL: {job.get('url', '')}",
            f"Location: {job.get('location', '')}",
            f"Pay: {job.get('pay', '')}",
            f"Work auth risk: {job.get('workAuthRisk', '')}",
            "",
            "Notes:",
            job.get("notes", ""),
            "",
            "Job description:",
            job.get("jd", ""),
        ]
    )
    out.write_text(content, encoding="utf-8")
    return out


def main():
    if len(sys.argv) < 3:
        raise SystemExit("Usage: generate_documents.py <resume|cover|all> <job_id>")
    kind = sys.argv[1]
    job = load_job(sys.argv[2])
    result = {"jobId": job["id"], "baseResumePath": str(base_resume_path())}
    if kind in ("resume", "all"):
        result["resumePath"] = str(make_resume(job))
    if kind in ("cover", "all"):
        result["coverPath"] = str(make_cover(job))
    result["jdPath"] = str(save_jd(job))
    print(json.dumps(result))


if __name__ == "__main__":
    main()

