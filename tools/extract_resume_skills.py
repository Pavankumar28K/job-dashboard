import json
import re
import sys
from pathlib import Path

from docx import Document


SKILL_TERMS = [
    "Python", "Java", "JavaScript", "TypeScript", "C#", ".NET", "ASP.NET Core",
    "Node.js", "React", "Angular", "HTML", "CSS", "SQL", "SQL Server",
    "PostgreSQL", "MySQL", "MongoDB", "Redis", "REST API", "GraphQL",
    "FastAPI", "Django", "Flask", "Spring Boot", "Express", "OOP",
    "Data Structures", "System Design", "Microservices", "Machine Learning",
    "Deep Learning", "NLP", "LLM", "RAG", "Prompt Engineering", "OpenAI",
    "LangChain", "Vector Databases", "Embeddings", "Pandas", "NumPy",
    "Scikit-learn", "PyTorch", "TensorFlow", "MLflow", "Statistics",
    "Data Analysis", "Data Modeling", "ETL", "Spark", "Airflow", "dbt",
    "Snowflake", "BigQuery", "Databricks", "Power BI", "Tableau", "Excel",
    "AWS", "Azure", "GCP", "Cloud", "Docker", "Kubernetes", "Terraform",
    "CI/CD", "Azure DevOps", "GitHub Actions", "Linux", "Bash", "Monitoring",
    "Selenium", "Playwright", "Cypress", "Testing", "API Testing",
    "Security", "OWASP", "SIEM", "Cloud Security", "Networking",
    "Salesforce", "CRM", "Apex", "SOQL", "Lightning", "Figma", "UX Design",
    "User Research", "Wireframing", "Prototyping", "Agile", "Scrum", "Jira",
    "Stakeholder Management", "Communication", "Documentation",
]


ALIASES = {
    "C#": ["c#", "c sharp", "csharp"],
    ".NET": [".net", "dotnet"],
    "ASP.NET Core": ["asp.net core", "asp.net", "aspnet"],
    "Node.js": ["node.js", "nodejs", "node js"],
    "REST API": ["rest api", "restful api", "rest apis", "api development"],
    "CI/CD": ["ci/cd", "cicd", "continuous integration", "continuous deployment"],
    "GitHub Actions": ["github actions"],
    "Azure DevOps": ["azure devops"],
    "Power BI": ["power bi", "powerbi"],
    "Scikit-learn": ["scikit-learn", "sklearn"],
    "Vector Databases": ["vector database", "vector databases", "vector db"],
}


def docx_text(path):
    doc = Document(path)
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def has_term(blob, term):
    terms = ALIASES.get(term, [term])
    for item in terms:
        item = item.lower()
        if re.search(rf"(?<![a-z0-9+#.]){re.escape(item)}(?![a-z0-9+#.])", blob):
            return True
    return False


def main():
    if len(sys.argv) < 2:
        raise SystemExit("Usage: extract_resume_skills.py <resume.docx> [extra skills json]")
    path = Path(sys.argv[1]).expanduser().resolve()
    if not path.is_file() or path.suffix.lower() != ".docx":
        raise SystemExit(f"Resume not found or not a .docx file: {path}")
    terms = list(SKILL_TERMS)
    if len(sys.argv) >= 3 and sys.argv[2].strip():
        try:
            for item in json.loads(sys.argv[2]):
                text = str(item).strip()
                if text and text not in terms:
                    terms.append(text)
        except json.JSONDecodeError:
            pass
    blob = docx_text(path).lower()
    skills = [term for term in terms if has_term(blob, term)]
    print(json.dumps({"resumePath": str(path), "skills": skills}))


if __name__ == "__main__":
    main()
